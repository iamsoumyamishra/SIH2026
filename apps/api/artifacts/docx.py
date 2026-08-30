"""DOCX artifact generation (AGENTS.md §22 — DOCX is the priority).

Creates real, valid .docx files using python-docx. The agent never returns
markdown pretending to be a document.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def generate_docx(
    path: str | Path,
    title: str,
    sections: list[dict],
) -> Path:
    """Create a .docx file.

    sections: list of {"heading": str, "body": str or [str]} entries.
    """
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)

    doc = DocxDocument()
    # Title
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for section in sections:
        heading = section.get("heading")
        body = section.get("body", "")
        if heading:
            doc.add_heading(heading, level=1)
        if isinstance(body, str):
            body = [body]
        for para_text in body:
            p = doc.add_paragraph()
            run = p.add_run(para_text)
            run.font.size = Pt(11)

    doc.save(str(path))
    return path


def make_approval_note(
    path: str | Path,
    machine_id: str,
    date: str,
    findings: list[dict],
    sop_references: list[str],
    recommendation: str,
    inconclusive: bool = False,
) -> Path:
    """Generate a standard approval note from inspection findings.

    `inconclusive=True` means no findable inspection items — the note must NOT
    fabricate an approval; it says review is required instead.
    """
    ok = [f for f in findings if f.get("status", "").upper() == "PASS"]
    failed = [f for f in findings if f.get("status", "").upper() == "FAIL"]

    if inconclusive:
        verdict = (
            "REVIEW REQUIRED — no inspection items were extracted from the uploaded "
            "document, so no automated decision can be issued. Confirm the source "
            "document is an inspection report and re-submit."
        )
    elif failed:
        verdict = (
            "CONDITIONALLY APPROVED — corrective maintenance on the failed items and "
            "re-inspection are required before return to service, per the referenced "
            "maintenance SOP."
        )
    else:
        verdict = "APPROVED for continued service."

    machine_label = machine_id or "Not identified in the document"
    title = f"Inspection Approval Note — {machine_id}" if machine_id else "Inspection Approval Note"

    sections: list[dict] = [
        {
            "heading": "Machine & Report",
            "body": [
                f"Machine ID: {machine_label}",
                f"Date of inspection: {date}",
                f"Items inspected: {len(findings)}  (Passed: {len(ok)}, Failed: {len(failed)})",
            ],
        },
        {
            "heading": "Findings",
            "body": [f"- {f.get('item', '')}: {f.get('remark', '')}" for f in findings],
        },
        {
            "heading": "Required Action",
            "body": recommendation,
        },
        {
            "heading": "SOP References",
            "body": sop_references or ["-"],
        },
        {
            "heading": "Approval Status",
            "body": verdict,
        },
    ]
    return generate_docx(path, title=title, sections=sections)
