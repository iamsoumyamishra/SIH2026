"""Document processing pipeline (AGENTS.md §17, §18).

Input → file-type detection → parse → scanned? → OCR → tables → normalized text
representation. Everything is local. OCR absence is reported explicitly and
never silently downgraded to a cloud service.
"""
from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from multimodal.images import load_image, render_pdf_pages
from multimodal.ocr import OcrEngine, OcrUnavailableError
from multimodal.pdf import extract_pdf_text, is_scanned
from multimodal.tables import extract_tables


@dataclass
class ExtractedDocument:
    filename: str
    content_type: str = "unknown"  # text | scanned | image
    text: str = ""
    pages: list[str] = field(default_factory=list)
    tables: list[dict[str, Any]] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)
    warnings: list[str] = field(default_factory=list)


class DocumentPipeline:
    def __init__(self, ocr_engine: OcrEngine | None = None, render_max_pages: int = 20) -> None:
        self.ocr = ocr_engine or OcrEngine()
        self.render_max_pages = render_max_pages

    def ingest(self, path: str | Path) -> ExtractedDocument:
        """Process a single file and return a normalized representation."""
        path = Path(path)
        suffix = path.suffix.lower()
        doc = ExtractedDocument(filename=path.name)

        if suffix in (".pdf",):
            return self._ingest_pdf(path, doc)
        if suffix in (".txt", ".md", ".log", ".csv"):
            return self._ingest_text(path, doc)
        if suffix in (".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".webp"):
            return self._ingest_image(path, doc)
        if suffix in (".docx",):
            return self._ingest_docx(path, doc)
        if suffix in (".xlsx", ".xls", ".csv"):
            return self._ingest_xlsx(path, doc)

        doc.warnings.append(f"Unsupported file type: {suffix or 'none'}")
        return doc

    # ── handlers ─────────────────────────────────────────────
    def _ingest_pdf(self, path: Path, doc: ExtractedDocument) -> ExtractedDocument:
        result = extract_pdf_text(path)
        doc.pages = result["pages"]
        doc.metadata["page_count"] = result["page_count"]
        doc.metadata["has_text_layer"] = result["has_text"]

        if not is_scanned(result):
            doc.content_type = "text"
            doc.text = result["full_text"]
        else:
            doc.content_type = "scanned"
            doc.warnings.append("Scanned PDF detected — running OCR.")
            try:
                if not self.ocr.is_available():
                    raise OcrUnavailableError("OCR engine unavailable")
                page_images = render_pdf_pages(path, max_pages=self.render_max_pages)
                doc.text = self.ocr.ocr_pdf_pages(page_images)
            except OcrUnavailableError as exc:
                doc.warnings.append(f"OCR unavailable: {exc}")
                doc.text = ""
            except Exception as exc:  # noqa: BLE001
                doc.warnings.append(f"OCR failed: {exc}")
                doc.text = ""

        # Tables (work for digital PDFs; best-effort for scanned)
        if doc.content_type == "text":
            doc.tables = extract_tables(path)
        return doc

    def _ingest_text(self, path: Path, doc: ExtractedDocument) -> ExtractedDocument:
        doc.content_type = "text"
        doc.text = path.read_text(encoding="utf-8", errors="replace")
        doc.pages = [doc.text]
        return doc

    def _ingest_image(self, path: Path, doc: ExtractedDocument) -> ExtractedDocument:
        doc.content_type = "image"
        try:
            if not self.ocr.is_available():
                raise OcrUnavailableError("OCR engine unavailable")
            img = load_image(path)
            doc.text = self.ocr.ocr_image(img)
        except OcrUnavailableError as exc:
            doc.warnings.append(f"OCR unavailable: {exc}")
            doc.text = ""
        except Exception as exc:  # noqa: BLE001
            doc.warnings.append(f"OCR failed: {exc}")
            doc.text = ""
        doc.pages = [doc.text]
        return doc

    def _ingest_docx(self, path: Path, doc: ExtractedDocument) -> ExtractedDocument:
        from docx import Document as DocxDocument

        d = DocxDocument(str(path))
        paragraphs = [p.text for p in d.paragraphs if p.text]
        tables = [cell.text for table in d.tables for row in table.rows for cell in row.cells]
        doc.content_type = "text"
        doc.text = "\n".join(paragraphs + tables)
        doc.pages = [doc.text]
        return doc

    def _ingest_xlsx(self, path: Path, doc: ExtractedDocument) -> ExtractedDocument:
        import openpyxl

        wb = openpyxl.load_workbook(str(path), data_only=True)
        lines: list[str] = []
        for ws in wb.worksheets:
            lines.append(f"# {ws.title}")
            for row in ws.iter_rows(values_only=True):
                values = ["" if c is None else str(c) for c in row]
                if any(v.strip() for v in values):
                    lines.append(", ".join(values))
        doc.content_type = "text"
        doc.text = "\n".join(lines)
        doc.pages = [doc.text]
        return doc
